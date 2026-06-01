"""Minimal Markdown -> .docx converter (python-docx) for the whitepaper.

Handles the subset used in docs/research/019: # / ## / ### headings, **bold**,
*italic*, `code`, pipe tables (with header + --- separator), - bullets,
N. numbered lists, > blockquotes, --- horizontal rules, and paragraphs.

  uv run python scripts/md2docx.py docs/research/019-associative-memory-whitepaper.md out.docx
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)")


def add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold**, *italic*, `code`."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> None:
    src, dst = sys.argv[1], sys.argv[2]
    with open(src, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # table block: a line starting with | followed by a |---| separator
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = split_row(line)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(re.sub(r"\*\*", "", h)); run.bold = True
            for r in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    cells[j].paragraphs[0].text = ""
                    add_runs(cells[j].paragraphs[0], r[j] if j < len(r) else "")
            doc.add_paragraph()
            continue

        if not stripped:
            i += 1; continue

        if stripped == "---":
            doc.add_paragraph().add_run("—" * 20).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            i += 1; continue

        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style="Title"); i += 1; continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1); i += 1; continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2); i += 1; continue
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote"); add_runs(p, stripped[2:]); i += 1; continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, stripped[2:]); i += 1; continue
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number"); add_runs(p, m.group(2)); i += 1; continue

        p = doc.add_paragraph(); add_runs(p, stripped); i += 1

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main()
