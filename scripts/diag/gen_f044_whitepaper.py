"""Generate the F044 tinyHippo-Lite algorithm whitepaper (.docx)."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def P(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def num(text):
    return doc.add_paragraph(text, style="List Number")

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    for line in text.split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(8.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

# ============================ TITLE ============================
title = doc.add_heading("F044 — tinyHippo-Lite", level=0)
sub = P("Algorithmic Sleep Consolidation for the Nous Memory Graph", bold=True, size=13)
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
meta = P("Technical Whitepaper  ·  Algorithm Flow, Activation, and Empirical Findings", italic=True, size=10)
P("Version 1.1  ·  2026-06-15  ·  Status: mechanism implemented & flag-gated (default OFF); unvalidated in single-pass usage", italic=True, size=9)
doc.add_paragraph()

# ============================ 1. EXEC SUMMARY ============================
H("1. Executive Summary", 1)
P("tinyHippo-Lite (F044) ports the biological Synaptic Tagging & Capture (STC) procedure — "
  "tag, reinforce, promote, homeostatically downscale — onto Nous's knowledge graph "
  "(brain.graph_edges) as a pure-SQL/Python algorithm. The goal: during offline 'sleep' "
  "cycles, let frequently-reactivated associations consolidate and relatively strengthen, so "
  "retrieval surfaces the durable, repeatedly-useful links rather than every cosine neighbour.")
P("This document specifies the full algorithm flow, then focuses on the question that determines "
  "whether F044 does anything in practice: activation. The mechanism is correct and spec-faithful, "
  "but it only produces an effect when a repeating reinforce → sleep → consolidate → downscale loop "
  "runs over time. A single-pass ingest never closes that loop, so F044 is inert there. The headline "
  "numbers in early harness tests (+5.1pp MRR) were produced by a manual warm-up that forced the loop "
  "artificially, and are withdrawn as a harness artifact. See §9–§10.")

# ============================ 2. THESIS ============================
H("2. Motivation & Biological Thesis", 1)
P("Nous's sleep cycle already does cosine-similarity graph densification (F040) and dead-edge "
  "pruning. What it lacked, and what F044 adds, is three STC ideas lifted from Talanov's "
  "hippocampal microcircuit work and the Tononi–Cirelli Synaptic Homeostasis Hypothesis (SHY):")
bullet("a two-tier edge state (provisional 'tagged' vs durable 'consolidated') gated by sustained "
       "reinforcement (the PRP / 'plasticity-related protein' threshold);", "Synaptic Tagging & Capture — ")
bullet("retrieval == reactivation: an edge exercised during recall is reinforced, mirroring "
       "sharp-wave-ripple replay;", "Reactivation — ")
bullet("net potentiation accumulated during 'wake' is globally renormalized during 'sleep' by "
       "multiplicatively downscaling un-consolidated synapses, so consolidated ones become "
       "relatively dominant.", "Homeostatic downscale (SHY) — ")
P("Crucially, F044 lifts only the parts that do not depend on neural topology: the constants, the "
  "procedural ordering, and the state machine. There is no spiking simulation.")

# ============================ 3. DATA MODEL ============================
H("3. Data Model — the Edge State Machine", 1)
P("Migration 061 adds three columns to brain.graph_edges:")
code("consolidation_state  TEXT  NOT NULL DEFAULT 'tagged'   -- 'tagged' | 'consolidated'\n"
     "ltp_count            INT   NOT NULL DEFAULT 0          -- cumulative reinforcement (PRP analog)\n"
     "last_ltp_at          TIMESTAMPTZ NULL                  -- last reinforcement time (debounce/telemetry)")
P("Column-level defaults make the feature inert until reinforcement begins: every existing edge is a "
  "'tagged' / ltp_count=0 no-op, so feature-OFF behaviour is byte-identical to pre-F044 main. "
  "A CHECK constrains the two-value state; a partial index on (agent_id, ltp_count) WHERE "
  "consolidation_state='tagged' keeps the promotion scan on the tagged frontier only.")
P("State transitions:", bold=True)
table(["Transition", "When", "Effect"],
      [["Created", "edge first inserted", "state='tagged', ltp_count=0"],
       ["Reinforced", "edge reactivated (recall) or re-derived (linker conflict)", "ltp_count += n"],
       ["Promotion (Phase 8c)", "sleep, if ltp_count ≥ PRP (default 3)", "state='consolidated' (sticky)"],
       ["Downscale (Phase 8d)", "sleep, if state='tagged'", "weight *= α (consolidated exempt)"],
       ["Mortality (F053)", "prune, weight below floor & tagged", "DELETE (deferred — not in v1)"]],
      widths=[1.6, 3.0, 2.2])

# ============================ 4. ALGORITHM FLOW ============================
H("4. Algorithm Flow", 1)
P("F044 is a four-stage pipeline. Stages 1–2 run continuously / at sleep; stages 3–4 run inside one "
  "sleep phase (_phase_stc_consolidation), after F040 densification and before F053 dead-edge prune.")

H("Stage 1 — Reactivation signal (two sources)", 2)
P("A) Recall-touch (the theoretically-faithful trigger). After every run_recall_pipeline call, the "
  "edges whose BOTH endpoints appear in the final co-retrieved result set (excluding 'contradicts') "
  "are buffered in a process-global Counter keyed (source_id, target_id, relation). The read path "
  "stays write-free — only an in-memory increment.", italic=False)
code("# nous/api/retrieval_pipeline.py :: _record_recall_reactivation\n"
     "SELECT source_id, target_id, relation FROM brain.graph_edges\n"
     " WHERE agent_id=:a AND relation<>'contradicts'\n"
     "   AND source_id = ANY(candidate_ids) AND target_id = ANY(candidate_ids)\n"
     "# -> record_recall_touches([(s,t,rel), ...])  (buffered, not written)")
P("B) Write re-derivation. The two LIVE similarity linkers (brain._auto_link decision↔decision; "
  "graph_linker.create_edge cross-type, excluding 'structural' provenance) increment ltp_count on "
  "their ON CONFLICT branch — a conflict means the same edge was independently rediscovered, i.e. "
  "reinforced. Deterministic sleep-time rebuilders are deliberately NOT routed here (they re-derive "
  "the same edges every cycle and would inflate uniformly).")

H("Stage 2 — Reinforcement flush (at sleep)", 2)
P("flush_recall_touches() applies the buffered touches to ltp_count (UPDATE ... SET ltp_count = "
  "ltp_count + :n, last_ltp_at = now()) and clears the buffer only after the writes succeed "
  "(rollback-safe). This is where 'wake' reactivations are committed to the durable counter.")

H("Stage 3 — Promotion gate (Phase 8c)", 2)
P("Idempotent, agent-scoped single UPDATE:")
code("UPDATE brain.graph_edges\n"
     "   SET consolidation_state='consolidated'\n"
     " WHERE agent_id=:agent AND consolidation_state='tagged'\n"
     "   AND ltp_count >= :prp                  -- PRP default 3")
P("An edge reactivated at least PRP times becomes permanently 'consolidated'. The state is sticky: "
  "once promoted, it is never downscaled by Phase 8d.")

H("Stage 4 — Homeostatic α-downscale (Phase 8d) — the retrieval mechanism", 2)
P("Promotion (8c) runs before downscale (8d) so freshly-promoted edges are not penalized this cycle. "
  "The downscale is the spec's actual way consolidation influences retrieval — a GLOBAL, persistent "
  "edge-weight change read by every weight-based graph consumer:")
code("# nous/brain/tinyhippo_lite.py :: homeostatic_downscale  (Phase 8d)\n"
     "UPDATE brain.graph_edges\n"
     "   SET weight = weight * :alpha           -- α default 0.75, validated band [0.50, 0.90]\n"
     " WHERE agent_id=:agent AND consolidation_state='tagged'   -- consolidated edges EXEMPT")
P("Over successive cycles, tagged edges decay (α, α², α³, …) while consolidated edges hold full "
  "weight, so the durable associations come to dominate spreading activation, neighbour scoring, the "
  "adjacency boost, and the rerank-by-score order.")

# ============================ 5. ACTIVATION ============================
H("5. Activation — How F044 Turns On (and why it usually doesn't)", 1)
P("This is the decisive section. F044 produces an effect only when its reinforce → consolidate → "
  "downscale loop actually closes. That requires sustained reactivation, which in turn requires:")
num("The full SleepHandler to run (it owns Phases 8c/8d). In the production app this fires on session "
    "idle timeout. NOTE: the BEAM ingest harness bypasses SleepHandler and calls GraphDensifier "
    "directly, so the F044 phases never run there.")
num("Reinforcement to accumulate to PRP (≥3) on the SAME edge. The two sources have different "
    "longitudinal requirements:")
bullet("requires repeated retrievals that co-activate the same edge, buffered across MULTIPLE sleep "
       "cycles (a query touches an edge; the touch flushes at the next sleep; three such cycles "
       "promote it). A single retrieval, or retrievals with no intervening sleep flush, contribute "
       "nothing durable.", "Recall-touch — ")
bullet("requires the SAME (source,target,relation) to be independently rediscovered ≥3 times by the "
       "live linkers — i.e. content recurrence over time (the same association keeps being inferred "
       "as new, overlapping material is learned).", "Write re-derivation — ")
P("The single-pass failure mode (empirically confirmed, §9).", bold=True)
P("In one ingest, every edge is created exactly once (no ON CONFLICT → write re-derivation = 0), and "
  "ingest performs no retrievals (recall-touch = 0). Therefore ltp_count stays 0 everywhere, nothing "
  "promotes, nothing consolidates, and the α-downscale (if it even runs) scales ALL edges uniformly "
  "— which changes no ranking. F044 is inert.")
P("Conclusion: F044 is a LONGITUDINAL feature. It can only demonstrate value across many sessions in "
  "which the same associations are repeatedly retrieved and/or re-derived, with sleep cycles between "
  "them to flush and promote. No single-pass evaluation (BEAM or otherwise) can exhibit this.")

# ============================ 6. RETRIEVAL EFFECT ============================
H("6. The Retrieval Effect", 1)
P("Once edges are consolidated, the α-downscale changes which neighbours win. Because tagged edges "
  "have decayed while consolidated edges retain full weight, graph expansion and the rerank-by-score "
  "step rank consolidated-connected candidates higher. The F051 harness scores with "
  "rerank_by_score=True, so a graph-expanded gold reached via a consolidated edge can climb into the "
  "top-K and improve MRR / nDCG — but only if that gold's bridge edge was among the consolidated set.")
P("A note on the superseded v1.1 substitute (B2). Before implementing the spec's α-downscale, an "
  "interim consumer multiplied a consolidated edge's contribution to the adjacency-boost degree by "
  "tinyhippo_consolidated_boost_factor (default 2.0). It was found dead (null/negative) on every "
  "instrument and, separately, contained a normalization bug (the factor cancelled through max-degree "
  "normalization). It is retained default-OFF for reference only; the spec mechanism is the α-downscale.")

# ============================ 7. CONFIG ============================
H("7. Configuration Reference", 1)
table(["Env var", "Default", "Meaning"],
      [["NOUS_TINYHIPPO_LITE_ENABLED", "false", "Master switch: state machine, reinforcement, telemetry phase"],
       ["NOUS_TINYHIPPO_RECALL_TOUCH_ENABLED", "true*", "Stage 1A recall-touch recording (*only when master ON)"],
       ["NOUS_TINYHIPPO_PRP_THRESHOLD", "3", "Reinforcements needed to promote tagged → consolidated"],
       ["NOUS_TINYHIPPO_DOWNSCALE_ENABLED", "false", "Phase 8d α-downscale (the retrieval mechanism)"],
       ["NOUS_TINYHIPPO_ALPHA", "0.75", "Per-cycle tagged-weight decay factor (band 0.50–0.90)"],
       ["NOUS_TINYHIPPO_CONSOLIDATED_BOOST_FACTOR", "2.0", "Superseded B2 substitute (dead; retained OFF)"]],
      widths=[2.9, 0.7, 3.2])
P("Default posture: master OFF → fully inert (byte-identical to pre-F044). Master ON alone → "
  "telemetry-only (promotion + counts, no weight change). Master ON + downscale ON → full mechanism.")

# ============================ 8. WORKED EXAMPLE ============================
H("8. End-to-End Worked Example (longitudinal)", 1)
P("Day 1: user asks about 'RRF tuning'. Recall co-retrieves fact F_rrf and decision D_rrf; the edge "
  "F_rrf→D_rrf is buffered. Session idles → sleep flush: ltp_count(F_rrf→D_rrf)=1. Still tagged.")
P("Days 2–3: the topic recurs; two more recalls touch the same edge, each flushed at its sleep. "
  "ltp_count reaches 3 → Phase 8c promotes F_rrf→D_rrf to 'consolidated'. Phase 8d downscales the "
  "thousands of one-off tagged edges by α, leaving this edge at full weight.")
P("Day 4: a SOURCE-flavoured query ('how did we validate hybrid ranking?') retrieves F_rrf; graph "
  "expansion now ranks D_rrf above the decayed alternatives → the durable answer surfaces. THIS is "
  "the intended payoff — and it took four sessions with intervening sleeps to materialize.")

# ============================ 9. EMPIRICAL ============================
H("9. Empirical Validation — the full arc", 1)
P("F044 was tested on three instruments. The investigation reversed direction twice as bugs and "
  "shortcuts were found; the final clean test is decisive.")
table(["Test", "Mechanism", "Result", "Verdict"],
      [["BEAM-100K, n=100, clean+permuted", "B2 adjacency boost", "−0.7pp best / −2.6pp generalize", "B2 null (also a math bug)"],
       ["Prod graph-targeted, n=22 (harness)", "spec α-downscale, α≈0.42", "+5.1pp MRR; random-control=0", "consolidation load-bearing…"],
       ["…paired significance", "—", "2–3 of 22 movers, t=1.20, p≈0.23", "…but NOT significant"],
       ["Clean full-cycle ingest", "real pipeline, F044 ON", "0 edges consolidated (0/25,756)", "INERT — +5pp was warm-up artifact"]],
      widths=[2.4, 1.7, 2.0, 1.9])
P("Why +5.1pp was withdrawn. The prod harness manually drove ~200 recall queries to force "
  "consolidation, then manually applied the downscale. The confound control was clean (exempting 287 "
  "RANDOM edges gave +0.000; exempting the 287 CONSOLIDATED edges gave +0.051 — so the consolidation "
  "was genuinely load-bearing, not global sharpening). But the clean full-cycle test then showed that "
  "in real ingestion NEITHER reinforcement source fires, so the 287 consolidations never occur. The "
  "mechanism works; its activation precondition is simply never met in single-pass usage.")

# ============================ 10. LIMITATIONS ============================
H("10. Limitations & Activation Requirements", 1)
bullet("Inert in single-pass usage. Demonstrating value requires a longitudinal eval — simulated "
       "repeated sessions/queries over time with intervening sleeps — which is a separate, larger "
       "project.", "")
bullet("BEAM ingest bypasses SleepHandler (runs densification directly), so Phases 8c/8d do not run "
       "in that harness; the production app runs them on idle timeout.", "")
bullet("Recall-touch is a process-global buffer (single-process scope); cross-worker touches in a "
       "multi-worker deploy are a known v1.1 limitation.", "")
bullet("Recall-touch uses co-retrieval as a proxy for the spec's literal edge-traversal signal.", "")
bullet("n=22 graph-targeted qrels are decision-only and LLM-asserted (reviewed_by='auto'); the "
       "generator's yield collapsed to ~0% on the current corpus, blocking a larger significance test.", "")

# ============================ 11. IMPLEMENTATION MAP ============================
H("11. Implementation Map", 1)
table(["Component", "Location"],
      [["State columns + index", "sql/migrations/061_f044_tinyhippo_stc.sql"],
       ["Core: gate, telemetry, increment, downscale, recall buffer/flush", "nous/brain/tinyhippo_lite.py"],
       ["Reinforcement hooks (write re-derivation)", "nous/brain/brain.py (_auto_link), nous/brain/graph_linker.py (create_edge)"],
       ["Recall-touch recorder", "nous/api/retrieval_pipeline.py (_record_recall_reactivation)"],
       ["Sleep Phases 8c+8d", "nous/handlers/sleep_handler.py (_phase_stc_consolidation)"],
       ["Config flags", "nous/config.py (tinyhippo_*)"],
       ["ORM columns", "nous/storage/models.py (GraphEdge)"],
       ["Eval harnesses", "scripts/diag/f044_prod_downscale_*.py, f044_beam_*"]],
      widths=[3.2, 3.6])

# ============================ 12. CONCLUSION ============================
H("12. Conclusion", 1)
P("F044 tinyHippo-Lite is a correctly-implemented, spec-faithful STC consolidation mechanism: a "
  "two-tier edge state machine, two reinforcement sources, a PRP promotion gate, and a homeostatic "
  "α-downscale that lets durable associations dominate retrieval. Controlled experiments confirm the "
  "effect, when present, is genuinely driven by which edges consolidate — not an artifact of the "
  "downscale itself. However, its activation is fundamentally longitudinal: it needs the same "
  "associations reactivated to the PRP threshold across multiple sleep cycles. In single-pass usage "
  "that precondition is never met, so the feature is inert, and the early +5.1pp figure is withdrawn "
  "as a manual-warm-up artifact. Recommended posture: keep all F044 surfaces committed and default-OFF; "
  "validate (or retire) only via a dedicated longitudinal evaluation.")

out = "docs/reviews/F044-tinyhippo-lite-whitepaper.docx"
doc.save(out)
print("wrote", out)
